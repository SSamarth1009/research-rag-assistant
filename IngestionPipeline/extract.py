import os
import fitz
import pdfplumber
from PIL import Image
from docx import Document as DocxDocument
from doclayout_yolo import YOLOv10
from schemas.schemas import Document, Page, Region, BoundingBox, KnowledgeObject, KnowledgeType


#Adding Model for Layout Detection
LAYOUT_MODEL = YOLOv10("models/doclayout_yolo_docstructbench_imgsz1024.pt") # *A fully trained neural network that is trainded on multiple PDF's


###########################################################################
#---------------------PDF EXTRACTION FUNCTIONS-----------------------------
###########################################################################
LABEL_MAPPING = {
    "title": KnowledgeType.TITLE,

    "plain text": KnowledgeType.PARAGRAPH,

    "table": KnowledgeType.TsABLE,
    "table_caption": KnowledgeType.TABLE_CAPTION,

    "figure": KnowledgeType.FIGURE,
    "figure_caption": KnowledgeType.FIGURE_CAPTION,

    "isolate_formula": KnowledgeType.FORMULA,
    "formula_caption": KnowledgeType.FORMULA_CAPTION,

    "header": KnowledgeType.HEADER,
    "footer": KnowledgeType.FOOTER,

    "list": KnowledgeType.LIST,
}


def detect_layout(fitz_page) -> list[Region]:
    """
    Detect document layout regions (paragraphs, tables,
    figures, captions, formulas, etc.)
    Returns:
        List[Region]
    """
    pix = fitz_page.get_pixmap(dpi=200) # *Convert the PDF page into an image.
    image = Image.frombytes(  # *Pillow provides a standard image representation (PIL.Image) that the YOLO model accepts as input for layout detection.
        "RGB",
        (pix.width, pix.height),
        pix.samples,
    )
    results = LAYOUT_MODEL.predict(image)
    regions = []

    if not results:
        return regions
    result = results[0]

    for box in result.boxes:
        cls_id = int(box.cls)
        label = result.names[cls_id]
        #print(result.names)
        try:
            region_type = LABEL_MAPPING.get(label, KnowledgeType.UNKNOWN)
        except ValueError:
            region_type = KnowledgeType.UNKNOWN

        print(f"{label:20} ---> {region_type.value}")

        x1, y1, x2, y2 = box.xyxy[0].tolist()

        regions.append(
            Region(
                type=region_type,
                confidence=float(box.conf),
                bbox=BoundingBox(    # *We use the BoundingBox class to store the coordinates of the detected region in the image.
                    x1=x1,
                    y1=y1,
                    x2=x2,
                    y2=y2,
                ),
            )
        )

    return regions

def image_bbox_to_pdf_rect(fitz_page, bbox: BoundingBox) -> fitz.Rect:
    """
    Convert a YOLO/image-space bounding box into PDF coordinates.
    
    Convert the bounding box detected by YOLO (image coordinates) into
    PDF coordinates so PyMuPDF can crop the correct region.
    """
    pdf_rect = fitz_page.rect
    pdf_width = pdf_rect.width
    pdf_height = pdf_rect.height

    pix = fitz_page.get_pixmap(matrix=fitz.Matrix(1700 / pdf_width, 2200 / pdf_height))
    rendered_width = pix.width
    rendered_height = pix.height

    scale_x = pdf_width / rendered_width
    scale_y = pdf_height / rendered_height

    return fitz.Rect(
        bbox.x1 * scale_x,
        bbox.y1 * scale_y,
        bbox.x2 * scale_x,
        bbox.y2 * scale_y,
    )


def crop_region(fitz_page, bbox, output_path):
    """
    Crop a layout region from a PDF page using the bounding box returned
    by DocLayout-YOLO.
    """

    # Render the page at the same DPI used for layout detection
    render_pix = fitz_page.get_pixmap(dpi=200)

    image_width = render_pix.width
    image_height = render_pix.height

    # pdf_width = fitz_page.rect.width
    # pdf_height = fitz_page.rect.height

    # scale_x = pdf_width / image_width
    # scale_y = pdf_height / image_height

    # x1 = bbox.x1
    # y1 = bbox.y1
    # x2 = bbox.x2
    # y2 = bbox.y2
    pdf_bbox = image_bbox_to_pdf_rect(fitz_page, bbox)
    print(f"PDF Page Size : {fitz_page.rect}")
    print(f"Rendered Size : {image_width} x {image_height}")
    print(f"YOLO BBox     : {bbox}")
    # Convert image coordinates -> PDF coordinates
    # rect = fitz.Rect(
    #     x1 * scale_x,
    #     y1 * scale_y,
    #     x2 * scale_x,
    #     y2 * scale_y,
    # )
    rect = pdf_bbox

    # Clip to page boundaries
    rect = rect & fitz_page.rect

    # Skip invalid boxes
    if rect.is_empty or rect.width <= 1 or rect.height <= 1:
        return False
    pix = fitz_page.get_pixmap(
        clip=rect,
        dpi=300
    )
    pix.save(output_path)
    return True

def extract_pdf_tables(plumber_page) -> list[KnowledgeObject]:
    """
    Extract tables from a PDF page.
    """
    tables = []
    extracted_tables = plumber_page.extract_tables()

    for table in extracted_tables:
        tables.append(
            KnowledgeObject(
                type=KnowledgeType.TABLE,
                content=table,
                metadata={}
            )
        )
    return tables

def extract_pdf_figures(fitz_page,document_name: str,regions: list[Region],) -> list[KnowledgeObject]:

    figures = []
    document_stem = os.path.splitext(document_name)[0]
    figure_dir = os.path.join("figures", document_stem)
    os.makedirs(figure_dir, exist_ok=True)

    for region in regions:
        if region.type != KnowledgeType.FIGURE:
            continue
        image_path = os.path.join(
            figure_dir,
            f"page_{fitz_page.number + 1}_figure_{len(figures)+1}.png",
        )
        success = crop_region(
            fitz_page,
            region.bbox,
            image_path,
        )
        if not success:
            continue

        figures.append(
            KnowledgeObject(
                type=KnowledgeType.FIGURE,
                bbox=region.bbox,
                content=None,
                metadata={
                    "document_name": document_name,
                    "page_number": fitz_page.number + 1,
                    "image_path": image_path,
                    "confidence": region.confidence,
                    "caption": None,
                    "ocr_text": None,
                },
            )
        )
    return figures

def extract_pdf_text_regions(
    fitz_page,
    regions: list[Region],
) -> list[KnowledgeObject]:

    text_types = {
        KnowledgeType.TITLE,
        KnowledgeType.PARAGRAPH,
        KnowledgeType.HEADER,
        KnowledgeType.FOOTER,
        KnowledgeType.LIST,
    }

    knowledge_objects = []

    for region in regions:

        if region.type not in text_types:
            continue

        pdf_rect = image_bbox_to_pdf_rect(
            fitz_page,
            region.bbox,
        )

        text = fitz_page.get_text(
            "text",
            clip=pdf_rect,
        ).strip()

        if not text:
            continue

        knowledge_objects.append(
            KnowledgeObject(
                type=region.type,
                bbox=region.bbox,
                content=text,
                metadata={
                    "page_number": fitz_page.number + 1,
                    "confidence": region.confidence,
                },
            )
        )

    return knowledge_objects


def extract_pdf(pdf_path: str) -> Document:

    pdf = fitz.open(pdf_path)
    plumber_pdf = pdfplumber.open(pdf_path)
    document = Document(
        document_name=os.path.basename(pdf_path),
        document_type="pdf",
    )

    for page_number, (fitz_page, plumber_page) in enumerate(
        zip(pdf, plumber_pdf.pages),
        start=1,
    ):
        page = Page(page_number=page_number)

        # ---------- Layout ----------
        regions = detect_layout(fitz_page)
        page.regions.extend(regions)

        # ---------- Text ----------
        page.knowledge_objects.extend(
        extract_pdf_text_regions(
            fitz_page,
            regions,
        ))

        # ---------- Tables ----------
        page.knowledge_objects.extend(
            extract_pdf_tables(plumber_page)
        )

        # ---------- Figures ----------
        page.knowledge_objects.extend(
            extract_pdf_figures(
                fitz_page,
                document.document_name,
                regions,
            )
        )
        document.pages.append(page)

        print(
            f"Page {page_number}: "
            f"Regions={len(page.regions)}, "
            f"Objects={len(page.knowledge_objects)}"
        )

    pdf.close()
    plumber_pdf.close()

    return document

###########################################################################
#---------------------DOCX EXTRACTION FUNCTIONS-----------------------------
###########################################################################

def extract_docx(docx_path: str) -> Document:

    doc = DocxDocument(docx_path)
    document = Document(
        document_name=os.path.basename(docx_path),
        document_type="docx",
    )

    page = Page(page_number=1)
    content = "\n".join(
        para.text
        for para in doc.paragraphs
        if para.text.strip()
    )

    page.knowledge_objects.append(
        KnowledgeObject(
            type=KnowledgeType.PARAGRAPH,
            content=content,
        )
    )

    document.pages.append(page)
    return document


###########################################################################
#---------------------TXT EXTRACTION FUNCTIONS-----------------------------
###########################################################################

def extract_txt(txt_path: str) -> Document:
    with open(txt_path, "r", encoding="utf-8") as f:
        content = f.read()

    document = Document(
        document_name=os.path.basename(txt_path),
        document_type="txt",
    )
    page = Page(page_number=1)

    page.knowledge_objects.append(
        KnowledgeObject(
            type=KnowledgeType.PARAGRAPH,
            content=content,
        )
    )
    document.pages.append(page)

    return document


def extract_document(file_path):
    """
    Detect file type and call the appropriate extractor.
    """
    extension = os.path.splitext(file_path)[1].lower()

    if extension == ".pdf":
        return extract_pdf(file_path)
    elif extension == ".docx":
        return extract_docx(file_path)
    elif extension == ".txt":
        return extract_txt(file_path)
    else:
        raise ValueError(f"Unsupported file type: {extension}")

if __name__ == "__main__":
    docs_folder = "docs"

    for file in os.listdir(docs_folder):
        path = os.path.join(docs_folder, file)

        document = extract_document(path)

        print("\n" + "=" * 60)

        print(document.document_name)
        print(document.document_type)
        print(len(document.pages))

        if document.pages:

            first_page = document.pages[0]

            print("\nKnowledge Objects")

            for obj in first_page.knowledge_objects:

                print(
                    obj.type.value,
                    str(obj.content)[:150],
                )


#python -m IngestionPipeline.extract